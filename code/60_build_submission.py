"""
Remaining submission-package documents: title/author page, cover letter, and the completed
STROBE-RECORD checklist. Uses the same docx helpers as 50_build_docx.py.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
import manuscript_text as M

from docx import Document                                    # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH                # noqa: E402
from docx.oxml import OxmlElement                            # noqa: E402
from docx.oxml.ns import qn                                  # noqa: E402
from docx.shared import Cm, Pt                               # noqa: E402

from paths import DOCS

CORRESPONDING_EMAIL = "yosefad1305@gmail.com"


def new_doc(line_numbers=False, spacing=1.15):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    st.paragraph_format.line_spacing = spacing
    st.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.5)
        if line_numbers:
            ln = OxmlElement("w:lnNumType")
            ln.set(qn("w:countBy"), "1")
            ln.set(qn("w:start"), "1")
            ln.set(qn("w:restart"), "continuous")
            ln.set(qn("w:distance"), "360")
            s._sectPr.append(ln)
    return doc


def para(doc, text="", bold=False, italic=False, size=12, justify=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def table(doc, header, rows, size=8, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for c, name in zip(t.rows[0].cells, header):
        c.text = ""
        r = c.paragraphs[0].add_run(str(name))
        r.bold = True
        r.font.size = Pt(size)
        c.paragraphs[0].paragraph_format.line_spacing = 1.0
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = ""
            r = c.paragraphs[0].add_run(str(v))
            r.font.size = Pt(size)
            c.paragraphs[0].paragraph_format.line_spacing = 1.0
    return t


# ------------------------------------------------------------------ title / author page
doc = new_doc()
para(doc, M.TITLE, bold=True, size=14)
para(doc)
para(doc, "Yosef Adiniaev\u00b9\u00b2*, Alon Gorenshtein\u00b2\u00b3*, Tohar M Timor\u00b9, "
     "Eyal Klang\u00b2\u2074, Alexander Geftler\u2075")
para(doc, "* Y.A. and A.G. contributed equally to this work.", size=10)
para(doc)
para(doc, "Affiliations", bold=True)
for i, aff in enumerate([
        "Faculty of Medicine, University of Debrecen, Debrecen, Hungary",
        "BRIDGE GenAI Lab, Beth Israel Deaconess Medical Center, Boston, MA, USA",
        "Department of Neurology, Beth Israel Deaconess Medical Center, Harvard Medical School, "
        "Boston, MA, USA",
        "Department of Radiology, Beth Israel Deaconess Medical Center, Harvard Medical School, "
        "Boston, MA, USA",
        "Department of Orthopedic Surgery, Soroka Medical Center, Beer Sheva, Israel"], 1):
    para(doc, f"{i}. {aff}", size=11)
para(doc)
para(doc, "Corresponding author", bold=True)
para(doc, "Yosef Adiniaev, BRIDGE GenAI Lab, Beth Israel Deaconess Medical Center, 330 Brookline "
     f"Avenue, Boston, MA 02215, USA. Email: {CORRESPONDING_EMAIL}", size=11)
para(doc)
para(doc, "ORCID iDs", bold=True)
for name, orcid in [("Yosef Adiniaev", "0009-0002-7678-633X"),
                    ("Alon Gorenshtein", "0009-0000-7542-8608"),
                    ("Tohar M Timor", "0009-0000-9842-1636"),
                    ("Eyal Klang", "0000-0002-4567-3108"),
                    ("Alexander Geftler", "0000-0002-5175-2013")]:
    para(doc, f"{name} {orcid}", size=11)
para(doc)
para(doc, "Running head", bold=True)
para(doc, M.RUNNING_HEAD, size=11)
para(doc, "Keywords", bold=True)
para(doc, M.KEYWORDS, size=11)
para(doc, "Article type", bold=True)
para(doc, "Original Article", size=11)
para(doc)
para(doc, "Author contributions", bold=True)
para(doc, "Y.A. contributed to conceptualisation, methodology, software, formal analysis, data "
     "curation, visualisation, and writing of the original draft. A.Gor. contributed to "
     "methodology, validation, and review and editing. T.M.T. contributed to interpretation and "
     "review and editing. E.K. contributed to supervision, methodology, and review and editing. "
     "A.Gef. contributed to clinical interpretation, supervision, and review and editing. All "
     "authors read and approved the final manuscript.", size=11, justify=True)
doc.save(os.path.join(DOCS, "Title_page_and_authors.docx"))
print("wrote Title_page_and_authors.docx")

# ------------------------------------------------------------------ cover letter
doc = new_doc()
para(doc, "8 August 2026", size=11)
para(doc)
para(doc, "Professor Laurent Poirel", size=11)
para(doc, "Editor-in-Chief", size=11)
para(doc, "European Journal of Clinical Microbiology & Infectious Diseases", size=11)
para(doc)
para(doc, "Dear Professor Poirel,", size=12)
para(doc)
for t in [
    f"We submit for your consideration our manuscript, \u201c{M.TITLE},\u201d as an Original "
    "Article.",

    "Culture-based descriptions of bone and joint infection are increasingly drawn from "
    "electronic health record databases. A surgical series knows which specimen came from the "
    "infected bone; a database does not, because microbiology tables record a laboratory workflow "
    "label rather than an anatomical site. We asked how much the resulting estimates depend on "
    "which labels are analysed.",

    "In 7697 code-defined orthopaedic-infection episodes in MIMIC-IV, we separated specimen "
    "labels into a source-specific tier that names a musculoskeletal structure or an implant "
    "procedure, and a generic tier of tissue, biopsy and undifferentiated foreign-body labels. "
    "Between these cohorts, no growth was 48.0% against 34.0% and the polymicrobial fraction "
    "10.9% against 42.1%. Two conventions that are almost never stated moved the numbers by a "
    "similar amount: counting an explicit report of mixed flora as polymicrobial raised the "
    "pooled fraction from 42.1% to 50.8%, and scoping the first-isolate rule within rather than "
    "across tiers changed the source-specific isolate count by 43%.",

    "We then asked whether the label itself explains any of this. Within the 487 episodes that "
    "supplied both kinds of specimen, an episode-stratified per-specimen model did not detect an "
    "association with the label (odds ratio 0.86; 95% CI, 0.64-1.15, with patient-clustered "
    "uncertainty), and a one-to-one matched subset agreed. So we report that these estimates depend heavily on which labels are "
    "analysed, and we explicitly do not claim to have decomposed that dependence into an effect "
    "of the specimen and an effect of who was sampled. We also tested, and did not confirm, the "
    "usual explanation that a high methicillin-resistance fraction in this database reflects "
    "intensive-care case mix.",

    "The contribution is methodological rather than a clinical benchmark, and we have written it "
    "that way. We do not offer a target no-growth rate, we do not infer a cause for culture "
    "negativity that these data cannot show, and we report the pooled and undeduplicated figures "
    "alongside the primary ones so that readers can see what each analytic choice does. We are "
    "also explicit about what the tiers do not fix: a source-specific label identifies the kind "
    "of structure sampled, but MIMIC-IV records no laterality or operative linkage, so it does "
    "not establish that the specimen came from the site carrying the diagnosis code. The findings "
    "come from one institution and one labelling system, and external replication is the "
    "necessary next step. The cost of the approach is precision: the source-specific tier holds "
    "885 evaluable specimens, and we mark estimates its size will not support rather than "
    "presenting them.",

    "All analysis code, the specimen and organism dictionaries, and the derived summary outputs "
    "are publicly archived, so every number is re-runnable on identical data. An earlier version "
    "of this work, using the pooled specimen definition we now argue against, is posted as a "
    "preprint on medRxiv (https://doi.org/10.64898/2026.07.09.26357616) and has not been "
    "published in a peer-reviewed journal. The manuscript is not under consideration elsewhere, "
    "all authors have approved the submission, and no author reports a competing interest.",

    "Thank you for considering our work.",
]:
    para(doc, t, size=12, justify=True)
para(doc)
para(doc, "Sincerely,", size=12)
para(doc, "Yosef Adiniaev, on behalf of all authors", size=12)
para(doc, CORRESPONDING_EMAIL, size=12)
doc.save(os.path.join(DOCS, "Cover_letter.docx"))
print("wrote Cover_letter.docx")

# ------------------------------------------------------------------ STROBE-RECORD checklist
doc = new_doc(spacing=1.0)
para(doc, "STROBE-RECORD checklist", bold=True, size=14)
para(doc, M.TITLE, italic=True, size=11)
para(doc, "Locations are given as manuscript section names rather than page and line numbers, "
     "because pagination depends on the file the journal renders. The manuscript file carries "
     "continuous line numbering throughout for reviewer reference.", size=10)
para(doc)

ITEMS = [
    ("1a", "Indicate the study design with a commonly used term in the title or the abstract",
     "Reported", "Title (\u201cmeasurement study\u201d); Abstract, Methods"),
    ("1b", "Provide in the abstract an informative and balanced summary",
     "Reported", "Abstract"),
    ("RECORD 1.1", "The type of data used should be specified",
     "Reported", "Abstract, Methods; Methods, Data source and reporting"),
    ("RECORD 1.2", "The population and the timeframe should be reported",
     "Reported", "Methods, Cohort; Methods, Statistical analysis (era representation)"),
    ("RECORD 1.3", "If linkage between databases was conducted, this should be stated",
     "Not applicable", "Single database; no linkage performed"),
    ("2", "Explain the scientific background and rationale", "Reported", "Introduction"),
    ("3", "State specific objectives, including any prespecified hypotheses",
     "Reported", "Introduction, final paragraph"),
    ("4", "Present key elements of study design early in the paper",
     "Reported", "Methods, Data source and reporting"),
    ("5", "Describe the setting, locations, and relevant dates",
     "Partially reported", "Methods, Data source and reporting; Methods, Statistical analysis. "
     "Calendar dates are not recoverable because MIMIC-IV applies a random per-patient date "
     "shift; the anchor-year group is reported instead"),
    ("6a", "Give the eligibility criteria, and the sources and methods of selection",
     "Reported", "Methods, Cohort; Online Resource 1, eTable 6"),
    ("RECORD 6.1", "Detail the methods of outcome, exposure, predictor and confounder "
     "identification with a list of codes and algorithms",
     "Reported", "Methods, Cohort; Methods, Specimen label tiers; Methods, Culture result "
     "status; Online Resource 1, eTables 1, 6, 7, 8"),
    ("RECORD 6.2", "Describe any validation of the codes or algorithms used",
     "Not conducted; audit only",
     "No formal validation study was performed. There is no blinded reference standard, and no "
     "sensitivity, specificity, predictive value or agreement statistic is reported, because no "
     "criterion standard for specimen provenance or infection status exists in this database. "
     "What is provided is an audit: the classification rules are stated in full (Methods, "
     "Culture result status; Online Resource 1, eTable 7) and a random sample of classified "
     "specimens is printed with the laboratory comment text against which each classification "
     "can be checked (Online Resource 1, eTable 9)"),
    ("RECORD 6.3", "If validation was conducted for this study, provide the results",
     "Not applicable", "No validation study was conducted; see item 6.2. The accounting of how "
     "many specimens each rule excluded is given in Results, Cohort and specimen accounting, and "
     "Online Resource 1, eTable 2"),
    ("7", "Clearly define all outcomes, exposures, predictors, and effect modifiers",
     "Reported", "Methods, Culture result status; Methods, Isolate deduplication; Methods, "
     "Microbiological methods and susceptibility interpretation"),
    ("RECORD 7.1", "A complete list of codes and algorithms used should be provided",
     "Reported", "Online Resource 1, eTables 1, 6, 7, 8; public code repository"),
    ("8", "For each variable, give sources of data and details of methods of assessment",
     "Reported", "Methods, Data source and reporting; Methods, Microbiological methods"),
    ("9", "Describe any efforts to address potential sources of bias",
     "Reported", "Methods, Specimen label tiers; Methods, Culture result status; Methods, "
     "Isolate deduplication; Discussion, limitations paragraph"),
    ("10", "Explain how the study size was arrived at",
     "Reported", "Methods, Cohort (all eligible episodes); Results, Cohort and specimen "
     "accounting"),
    ("11", "Explain how quantitative variables were handled in the analyses",
     "Reported", "Methods, Statistical analysis"),
    ("12a", "Describe all statistical methods, including those used to control for confounding",
     "Reported", "Methods, Statistical analysis"),
    ("12b", "Describe any methods used to examine subgroups and interactions",
     "Reported", "Methods, Statistical analysis; Results, Within episodes, no per-specimen "
     "effect of the label is demonstrable"),
    ("12c", "Explain how missing data were addressed",
     "Reported", "Methods, Culture result status (indeterminate and cancelled tests excluded and "
     "counted); Online Resource 1, eTable 2"),
    ("12d", "If applicable, explain how loss to follow-up was addressed",
     "Not applicable", "Cross-sectional analysis of index admissions; no follow-up"),
    ("12e", "Describe any sensitivity analyses",
     "Reported", "Methods, Specimen label tiers; Methods, Isolate deduplication; Results, "
     "throughout; Tables 2-4; Online Resource 1, eTables 3-5"),
    ("RECORD 12.1", "Authors should describe the extent to which the investigators had access to "
     "the database population used to create the study population",
     "Reported", "Methods, Data source and reporting; Declarations, Data availability"),
    ("RECORD 12.2", "Authors should provide information on the data cleaning methods",
     "Reported", "Methods, Culture result status; Methods, Isolate deduplication; Online "
     "Resource 1, eTables 7-9"),
    ("RECORD 12.3", "State whether the study included person-level, institutional-level, or other "
     "data linkage across two or more databases",
     "Not applicable", "No linkage; a single database was used"),
    ("13", "Report numbers of individuals at each stage of the study",
     "Reported", "Results, Cohort and specimen accounting; Table 1; Online Resource 1, eTable 2"),
    ("14", "Give characteristics of study participants and information on exposures and "
     "potential confounders", "Reported", "Table 1"),
    ("15", "Report numbers of outcome events or summary measures",
     "Reported", "Results, all subsections; Tables 2-4; Figs. 1-3"),
    ("16", "Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their "
     "precision", "Reported", "Results, Exploratory infection-type contrast "
     "(models with and without adjustment for specimen count); Results, Within episodes, no "
     "per-specimen effect of the label is demonstrable (episode-stratified conditional model and "
     "one-to-one matched subset); all proportions carry "
     "patient-clustered 95% CIs"),
    ("17", "Report other analyses done", "Reported",
     "Results, Resistance and the effect of repeated isolates (era and intensive-care "
     "stratification); Online Resource 1, eTables 3-5"),
    ("RECORD 13.1", "Describe in detail the selection of the persons included in the study and "
     "the reasons for exclusion", "Reported",
     "Methods, Cohort; Results, Cohort and specimen accounting"),
    ("18", "Summarise key results with reference to study objectives",
     "Reported", "Discussion, first paragraph"),
    ("19", "Discuss limitations, taking into account sources of potential bias or imprecision",
     "Reported", "Discussion, limitations paragraph"),
    ("20", "Give a cautious overall interpretation of results",
     "Reported", "Discussion, third and fourth paragraphs"),
    ("21", "Discuss the generalisability of the study results",
     "Reported", "Discussion, limitations paragraph"),
    ("RECORD 19.1", "Discuss the implications of using data that were not created or collected to "
     "answer the specific research question", "Reported",
     "Introduction, second paragraph; Methods, Specimen label tiers; Discussion, second and "
     "fourth paragraphs"),
    ("22", "Give the source of funding and the role of the funders",
     "Reported", "Declarations, Funding"),
    ("RECORD 22.1", "Authors should provide information on how to access any supplemental "
     "information", "Reported", "Declarations, Data availability and Code availability"),
]
table(doc, ["Item", "Recommendation", "Status", "Location in manuscript"],
      [[a, b, c, d] for a, b, c, d in ITEMS], size=8)
doc.save(os.path.join(DOCS, "STROBE-RECORD_checklist.docx"))
print("wrote STROBE-RECORD_checklist.docx")
