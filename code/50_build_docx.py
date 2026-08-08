"""
Assemble the submission Word documents.

Journal formatting requirements handled here:
  * continuous line numbering across the whole manuscript (sectPr/lnNumType);
  * plain-text bracketed citations and a plain-text reference list, with no reference-manager
    field codes anywhere in the file;
  * double spacing, A4, 2.5 cm margins, page numbers.
"""
import json
import os
import platform
import sys

import matplotlib
import numpy as np
import pandas as pd
import scipy
import statsmodels
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, os.path.dirname(__file__))
import manuscript_text as M
from org_map import classify as org_classify

from paths import OUT, FIG, DOCS

with open(os.path.join(OUT, "tables.json"), encoding="utf-8") as f:
    T = json.load(f)


# ------------------------------------------------------------------ docx helpers
def new_doc(line_numbers=True, double_space=True):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = st.paragraph_format
    pf.line_spacing = 2.0 if double_space else 1.15
    pf.space_after = Pt(0)

    for s in doc.sections:
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.5)
        if line_numbers:
            add_line_numbers(s)
    add_page_numbers(doc)
    return doc


def add_line_numbers(section, start=1, distance_twips=360):
    """Continuous line numbering across the section, as the journal requires."""
    sect_pr = section._sectPr
    ln = sect_pr.find(qn("w:lnNumType"))
    if ln is None:
        ln = OxmlElement("w:lnNumType")
        sect_pr.append(ln)
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), str(start))
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), str(distance_twips))


def add_page_numbers(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        for instr in ("begin", "PAGE", "end"):
            el = OxmlElement("w:fldChar" if instr != "PAGE" else "w:instrText")
            if instr == "PAGE":
                el.set(qn("xml:space"), "preserve")
                el.text = " PAGE "
            else:
                el.set(qn("w:fldCharType"), instr)
            run._r.append(el)


def h(doc, text, size=13, after=6, before=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def para(doc, text, italic=False, justify=False, size=12, spacing=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if spacing:
        p.paragraph_format.line_spacing = spacing
    return p


def add_table(doc, spec, font_size=9):
    para(doc, spec["caption"], size=10, spacing=1.0)
    t = doc.add_table(rows=1, cols=len(spec["header"]))
    t.style = "Table Grid"
    for c, name in zip(t.rows[0].cells, spec["header"]):
        c.text = ""
        r = c.paragraphs[0].add_run(str(name))
        r.bold = True
        r.font.size = Pt(font_size)
        c.paragraphs[0].paragraph_format.line_spacing = 1.0
    for row in spec["rows"]:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = ""
            r = c.paragraphs[0].add_run(str(v))
            r.font.size = Pt(font_size)
            c.paragraphs[0].paragraph_format.line_spacing = 1.0
    doc.add_paragraph()
    return t


# ------------------------------------------------------------------ manuscript
doc = new_doc()

para(doc, M.TITLE, size=14).runs[0].bold = True
para(doc, f"Running head: {M.RUNNING_HEAD}", size=11)
doc.add_paragraph()

h(doc, "Abstract", before=0)
for label, body in M.ABSTRACT:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}. ")
    r.bold = True
    p.add_run(body)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Keywords ").bold = True
p.add_run(M.KEYWORDS)

doc.add_page_break()
h(doc, "Introduction", before=0)
for t in M.INTRODUCTION:
    para(doc, t, justify=True)

h(doc, "Materials and Methods")
for title, paras in M.METHODS:
    h(doc, title, size=12, before=8, after=2)
    for t in paras:
        para(doc, t, justify=True)

h(doc, "Results")
for title, paras in M.RESULTS:
    h(doc, title, size=12, before=8, after=2)
    for t in paras:
        para(doc, t, justify=True)

h(doc, "Discussion")
for t in M.DISCUSSION:
    para(doc, t, justify=True)

doc.add_page_break()
h(doc, "References", before=0)
for i, ref in enumerate(M.REFERENCES, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.add_run(f"[{i}] {ref}")

doc.add_page_break()
h(doc, "Statements and Declarations", before=0)
for label, body in M.DECLARATIONS:
    p = doc.add_paragraph()
    p.add_run(f"{label}. ").bold = True
    p.add_run(body)

doc.add_page_break()
h(doc, "Tables", before=0)
for key, num in [("table1", 1), ("table2", 2), ("table3", 3), ("table4", 4)]:
    spec = dict(T[key])
    spec["caption"] = f"Table {num}. {spec['caption']}"
    add_table(doc, spec)

doc.add_page_break()
h(doc, "Figures", before=0)
# Figures are placed in the manuscript, as the journal prefers, and are also supplied as separate
# 600 dpi PNG and vector PDF files at the journal's 174 mm production width. In this A4 document
# they are inserted at 16.0 cm, the usable text width between 2.5 cm margins, so nothing overruns.
for (tag, body), stem in zip(M.FIGURE_LEGENDS, ["Fig1", "Fig2", "Fig3"]):
    png = os.path.join(FIG, f"{stem}.png")
    if os.path.exists(png):
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.line_spacing = 1.0
        pic.add_run().add_picture(png, width=Cm(16.0))
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.add_run(f"{tag} ").bold = True
    p.add_run(body)
    if stem != "Fig3":
        doc.add_page_break()

path = os.path.join(DOCS, "manuscript.docx")
doc.save(path)
print("wrote", path)

# ------------------------------------------------------------------ supplement
sup = new_doc(line_numbers=False, double_space=False)
para(doc=sup, text="Online Resource 1", size=16).runs[0].bold = True
para(sup, "Supplementary material for:", size=11)
para(sup, M.TITLE, size=12, italic=True)
para(sup, "European Journal of Clinical Microbiology & Infectious Diseases", size=11)
sup.add_paragraph()
para(sup, "Yosef Adiniaev, Alon Gorenshtein, Tohar M Timor, Eyal Klang, Alexander Geftler",
     size=11)
para(sup, "Corresponding author: Yosef Adiniaev, BRIDGE GenAI Lab, Beth Israel Deaconess "
     "Medical Center, Boston, MA, USA; and Faculty of Medicine, University of Debrecen, "
     "Debrecen, Hungary.", size=11)
para(sup, "Correspondence: yosefad1305@gmail.com", size=11)
sup.add_page_break()

h(sup, "Contents", before=0)
etables = [("etable_specimen_map", "eTable 1", "Specimen label to tier assignment"),
           ("etable_accounting", "eTable 2", "Specimen accounting for the no-growth denominator"),
           ("etable_poly", "eTable 3", "Polymicrobial fraction under each prespecified rule"),
           ("etable_gradient", "eTable 4", "No growth by number of source-specific specimens per episode"),
           ("etable_era", "eTable 5", "Resistance and no growth by anchor-year group")]
# Contents lists every eTable in the document, including those built further down from sources
# other than tables.json. Keep this list in step with the tables actually written below.
CONTENTS = [(tag, title) for _, tag, title in etables] + [
    ("eTable 6", "Diagnosis codes defining the cohort"),
    ("eTable 7", "Rules assigning a culture result status"),
    ("eTable 8", "Organism normalisation dictionary"),
    ("eTable 9", "Software environment"),
    ("eTable 10", "Complete patient-clustered logistic model output"),
    ("eTable 11", "Exact versus patient-clustered confidence intervals"),
    ("eTable 12", "Within-episode comparison of specimen label tiers (primary comparison)"),
]
for tag, title in CONTENTS:
    para(sup, f"{tag}. {title}", size=11)
sup.add_page_break()

for key, tag, _ in etables:
    spec = dict(T[key])
    spec["caption"] = f"{tag}. {spec['caption']}"
    add_table(sup, spec, font_size=8)

# ---- eTable 6: diagnosis code list
sup.add_page_break()
add_table(sup, dict(
    caption=("eTable 6. Diagnosis codes defining the cohort. Codes are matched as prefixes in "
             "any diagnosis position."),
    header=["Group", "ICD-9-CM", "ICD-10-CM"],
    rows=[["Prosthetic joint infection", "996.66", "T84.5x"],
          ["Other internal orthopaedic device infection", "996.67", "T84.6x, T84.7x"],
          ["Native osteomyelitis", "730.xx", "M86.xx"]]), font_size=9)

# ---- eTable 7: culture result-status rules
add_table(sup, dict(
    caption=("eTable 7. Rules assigning a culture result status. Only positive and reported "
             "no-growth specimens are evaluable for the no-growth analysis."),
    header=["Status", "Rule", "Disposition"],
    rows=[["Positive", "At least one recognised organism on a bacterial-culture row",
           "Evaluable, numerator of culture-positive"],
          ["Reported no growth",
           "No organism, and a bacterial-culture comment matching completed-negative report "
           "language (for example, no growth; no anaerobes isolated; no significant growth)",
           "Evaluable, numerator of no growth"],
          ["Cancelled",
           "Comment indicates cancellation, patient crediting, an unsuitable or rejected "
           "specimen, or a test not performed; or the organism field carries an administrative "
           "string", "Excluded"],
          ["Indeterminate",
           "No organism and no interpretable comment: empty, the de-identification placeholder, "
           "or language that does not establish a completed negative report (for example, an "
           "abbreviated workup because of overgrowth)", "Excluded"]]), font_size=8)

# ---- eTable 8: organism dictionary
sup.add_page_break()
odict = sorted({(v, org_classify(v).get("genus_group"), org_classify(v).get("broad_group"))
                for v in pd.read_parquet(
                    os.path.join(OUT, "intermediate", "organisms.parquet")).org_name.unique()})
add_table(sup, dict(
    caption=("eTable 8. Organism normalisation dictionary: every laboratory organism string "
             "observed in the cohort, with its assigned reporting group."),
    header=["Laboratory organism string", "Reporting group", "Broad group"],
    rows=[[a, str(b), str(c)] for a, b, c in odict]), font_size=7)


# ---- eTable 10: software versions
sup.add_page_break()
add_table(sup, dict(
    caption="eTable 9. Software environment. Analyses used fixed random seeds throughout.",
    header=["Component", "Version"],
    rows=[["Python", platform.python_version()],
          ["pandas", pd.__version__], ["numpy", np.__version__],
          ["scipy", scipy.__version__], ["statsmodels", statsmodels.__version__],
          ["matplotlib", matplotlib.__version__],
          ["Random seed", "20260808"],
          ["Bootstrap resamples", "2000 (clusters = patients)"]]), font_size=9)

# ---- eTables 11-12: model output and interval comparison
sup.add_page_break()
for key, tag in [("etable_regression", "eTable 10"), ("etable_exact", "eTable 11"),
                 ("etable_within", "eTable 12")]:
    spec = dict(T[key])
    if not spec["caption"].startswith("eTable"):
        spec["caption"] = f"{tag}. {spec['caption']}"
    add_table(sup, spec, font_size=8)

path = os.path.join(DOCS, "Online_Resource_1.docx")
sup.save(path)
print("wrote", path)
